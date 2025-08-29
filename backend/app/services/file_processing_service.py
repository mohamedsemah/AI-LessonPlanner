import os
import io
import base64
import logging
from typing import List, Dict, Any, Optional
import fitz  # PyMuPDF for PDF processing
from docx import Document
from PIL import Image

# Import pytesseract for optimal OCR
try:
    import pytesseract
    TESSERACT_AVAILABLE = True
    # Configure Tesseract for optimal text extraction
    # You can customize these settings based on your needs
    TESSERACT_CONFIG = '--oem 3 --psm 6 -c tessedit_char_whitelist=ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz0123456789.,!?@#$%^&*()_+-=[]{}|;:,.<>?/~` '
except ImportError:
    TESSERACT_AVAILABLE = False
    logging.warning("pytesseract not available. OCR functionality will be limited.")

import json

logger = logging.getLogger(__name__)

class FileProcessingService:
    def __init__(self):
        self.logger = logging.getLogger(__name__)
        
        # Try to configure Tesseract path if available
        if TESSERACT_AVAILABLE:
            try:
                # Common Tesseract installation paths on Windows
                possible_paths = [
                    r'C:\Program Files\Tesseract-OCR\tesseract.exe',
                    r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
                    r'C:\Users\moham\AppData\Local\Tesseract-OCR\tesseract.exe'
                ]
                
                for path in possible_paths:
                    if os.path.exists(path):
                        pytesseract.pytesseract.tesseract_cmd = path
                        self.logger.info(f"Tesseract configured at: {path}")
                        break
                else:
                    # If no path found, let pytesseract try to find it automatically
                    self.logger.info("Tesseract path not found, using system PATH")
                    
            except Exception as e:
                self.logger.warning(f"Could not configure Tesseract path: {e}")
    
    async def process_uploaded_files(self, uploaded_files: List[Dict[str, Any]]) -> Dict[str, Any]:
        """
        Process uploaded files and extract their content for AI analysis
        
        Args:
            uploaded_files: List of file objects from frontend
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        try:
            self.logger.info(f"Processing {len(uploaded_files)} uploaded files")
            
            processed_content = {
                "syllabus_content": "",
                "materials_content": "",
                "images_description": "",
                "data_insights": "",
                "file_metadata": [],
                "total_content_length": 0
            }
            
            for i, file_data in enumerate(uploaded_files):
                try:
                    file_info = await self._process_single_file(file_data, i)
                    if file_info:
                        processed_content["file_metadata"].append(file_info)
                        
                        # Categorize content based on file type and content
                        if file_info["file_type"] in ["pdf", "docx", "txt"]:
                            if "syllabus" in file_info["filename"].lower() or "course" in file_info["filename"].lower():
                                processed_content["syllabus_content"] += f"\n\n{file_info['filename']}:\n{file_info['extracted_text']}"
                            else:
                                processed_content["materials_content"] += f"\n\n{file_info['filename']}:\n{file_info['extracted_text']}"
                        
                        elif file_info["file_type"] in ["image"]:
                            processed_content["images_description"] += f"\n\n{file_info['filename']}:\n{file_info['image_description']}"
                            # Also add extracted text if available
                            if file_info.get("extracted_text"):
                                processed_content["materials_content"] += f"\n\n{file_info['filename']} (OCR):\n{file_info['extracted_text']}"
                        
                        processed_content["total_content_length"] += len(file_info.get("extracted_text", ""))
                        
                except Exception as e:
                    self.logger.error(f"Error processing file {i}: {str(e)}")
                    processed_content["file_metadata"].append({
                        "filename": file_data.get("name", f"file_{i}"),
                        "error": f"Failed to process: {str(e)}",
                        "file_type": "unknown"
                    })
            
            # Generate contextual summary for AI
            processed_content["ai_context"] = self._generate_ai_context(processed_content)
            
            self.logger.info(f"Successfully processed {len(processed_content['file_metadata'])} files")
            return processed_content
            
        except Exception as e:
            self.logger.error(f"Error in process_uploaded_files: {str(e)}")
            raise
    
    async def _process_single_file(self, file_data: Dict[str, Any], index: int) -> Optional[Dict[str, Any]]:
        """Process a single uploaded file"""
        try:
            filename = file_data.get("name", f"file_{index}")
            file_type = file_data.get("type", "")
            file_size = file_data.get("size", 0)
            
            self.logger.info(f"Processing file: {filename} (type: {file_type}, size: {file_size})")
            
            # Extract base64 content
            content = file_data.get("content", "")
            if not content:
                raise ValueError("No file content provided")
            
            # Remove data URL prefix if present
            if content.startswith("data:"):
                content = content.split(",", 1)[1]
            
            # Decode base64 content
            file_bytes = base64.b64decode(content)
            
            # Process based on file type
            if file_type == "application/pdf":
                return await self._process_pdf(file_bytes, filename, file_size)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self._process_docx(file_bytes, filename, file_size)
            elif file_type == "text/plain":
                return await self._process_txt(file_bytes, filename, file_size)
            elif file_type.startswith("image/"):
                return await self._process_image(file_bytes, filename, file_size)
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            self.logger.error(f"Error processing file {index}: {str(e)}")
            return None
    
    async def _process_pdf(self, file_bytes: bytes, filename: str, file_size: int) -> Dict[str, Any]:
        """Extract text content from PDF files"""
        try:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            text_content = ""
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                text_content += page.get_text()
            
            doc.close()
            
            return {
                "filename": filename,
                "file_type": "pdf",
                "file_size": file_size,
                "extracted_text": text_content.strip(),
                "pages": len(doc),
                "content_summary": self._summarize_content(text_content)
            }
            
        except Exception as e:
            self.logger.error(f"Error processing PDF {filename}: {str(e)}")
            raise
    
    async def _process_docx(self, file_bytes: bytes, filename: str, file_size: int) -> Dict[str, Any]:
        """Extract text content from DOCX files"""
        try:
            doc = Document(io.BytesIO(file_bytes))
            text_content = ""
            
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text_content += cell.text + "\t"
                    text_content += "\n"
            
            return {
                "filename": filename,
                "file_type": "docx",
                "file_size": file_size,
                "extracted_text": text_content.strip(),
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "content_summary": self._summarize_content(text_content)
            }
            
        except Exception as e:
            self.logger.error(f"Error processing DOCX {filename}: {str(e)}")
            raise
    
    async def _process_txt(self, file_bytes: bytes, filename: str, file_size: int) -> Dict[str, Any]:
        """Extract text content from TXT files"""
        try:
            text_content = file_bytes.decode('utf-8', errors='ignore')
            
            return {
                "filename": filename,
                "file_type": "txt",
                "file_size": file_size,
                "extracted_text": text_content.strip(),
                "lines": len(text_content.split('\n')),
                "content_summary": self._summarize_content(text_content)
            }
            
        except Exception as e:
            self.logger.error(f"Error processing TXT {filename}: {str(e)}")
            raise
    
    async def _process_image(self, file_bytes: bytes, filename: str, file_size: int) -> Dict[str, Any]:
        """Extract text and describe image content with optimal OCR settings"""
        try:
            image = Image.open(io.BytesIO(file_bytes))
            
            # Extract text using OCR with optimal settings
            extracted_text = ""
            ocr_success = False
            
            if TESSERACT_AVAILABLE:
                try:
                    # Preprocess image for better OCR results
                    processed_image = self._preprocess_image_for_ocr(image)
                    
                    # Try with optimal settings first
                    extracted_text = pytesseract.image_to_string(processed_image, config=TESSERACT_CONFIG)
                    if extracted_text.strip():
                        ocr_success = True
                        self.logger.info(f"OCR successful for {filename} with optimal settings")
                    else:
                        # Fallback to different PSM modes if no text found
                        fallback_configs = [
                            '--oem 3 --psm 3',  # Fully automatic page segmentation
                            '--oem 3 --psm 4',  # Assume a single column of text
                            '--oem 3 --psm 8',  # Single word
                            '--oem 3 --psm 11'  # Sparse text with OSD
                        ]
                        
                        for config in fallback_configs:
                            try:
                                fallback_text = pytesseract.image_to_string(processed_image, config=config)
                                if fallback_text.strip() and len(fallback_text.strip()) > len(extracted_text.strip()):
                                    extracted_text = fallback_text
                                    ocr_success = True
                                    self.logger.info(f"OCR successful for {filename} with fallback config: {config}")
                                    break
                            except Exception:
                                continue
                                
                except Exception as ocr_error:
                    # Silent fallback - don't log OCR errors to avoid spam
                    self.logger.debug(f"OCR failed for {filename}: {str(ocr_error)}")
                    extracted_text = ""
            
            # Generate image description
            image_description = self._describe_image(image, filename)
            
            # If OCR failed but we have a filename hint, try to extract meaningful info
            if not ocr_success and extracted_text.strip():
                # Clean up OCR text
                extracted_text = self._clean_ocr_text(extracted_text)
            
            return {
                "filename": filename,
                "file_type": "image",
                "file_size": file_size,
                "extracted_text": extracted_text.strip(),
                "image_description": image_description,
                "ocr_success": ocr_success,
                "dimensions": f"{image.width}x{image.height}",
                "format": image.format,
                "mode": image.mode,
                "content_summary": self._summarize_content(extracted_text) if extracted_text else image_description
            }
            
        except Exception as e:
            self.logger.error(f"Error processing image {filename}: {str(e)}")
            raise
    
    def _summarize_content(self, text: str, max_length: int = 500) -> str:
        """Generate a summary of the content"""
        if not text:
            return ""
        
        # Simple summary: first few sentences
        sentences = text.split('.')
        summary = '. '.join(sentences[:3]) + '.'
        
        if len(summary) > max_length:
            summary = summary[:max_length] + "..."
        
        return summary
    
    def _describe_image(self, image: Image.Image, filename: str) -> str:
        """Generate a description of the image"""
        try:
            # Basic image analysis
            width, height = image.size
            format_name = image.format or "Unknown"
            mode = image.mode
            
            # Analyze image characteristics
            if mode == "RGB":
                color_info = "Color image"
            elif mode == "L":
                color_info = "Grayscale image"
            else:
                color_info = f"{mode} mode image"
            
            # Check if it might be a graph, chart, or diagram
            if any(keyword in filename.lower() for keyword in ["graph", "chart", "diagram", "plot", "figure"]):
                content_type = "Likely contains visual data, graphs, or diagrams"
            elif any(keyword in filename.lower() for keyword in ["photo", "image", "picture"]):
                content_type = "Photograph or visual content"
            else:
                content_type = "Visual content"
            
            description = f"{content_type} ({color_info}, {width}x{height} pixels, {format_name} format)"
            
            return description
            
        except Exception as e:
            self.logger.warning(f"Error describing image {filename}: {str(e)}")
            return f"Image file ({image.format or 'Unknown'} format)"
    
    def _generate_ai_context(self, processed_content: Dict[str, Any]) -> str:
        """Generate contextual information for AI prompts with enhanced OCR content"""
        context_parts = []
        
        if processed_content["syllabus_content"]:
            context_parts.append("COURSE SYLLABUS CONTENT:")
            context_parts.append(processed_content["syllabus_content"][:2000] + "...")
        
        if processed_content["materials_content"]:
            context_parts.append("\n\nRELEVANT MATERIALS:")
            context_parts.append(processed_content["materials_content"][:1500] + "...")
        
        if processed_content["images_description"]:
            context_parts.append("\n\nVISUAL CONTENT:")
            context_parts.append(processed_content["images_description"][:1000] + "...")
            
            # Add OCR insights if available
            ocr_insights = []
            for file_info in processed_content.get("file_metadata", []):
                if file_info.get("file_type") == "image" and file_info.get("ocr_success"):
                    ocr_insights.append(f"• {file_info['filename']}: OCR extracted {len(file_info['extracted_text'])} characters")
            
            if ocr_insights:
                context_parts.append("\n\nOCR EXTRACTION RESULTS:")
                context_parts.extend(ocr_insights)
        
        if processed_content["data_insights"]:
            context_parts.append("\n\nDATA INSIGHTS:")
            context_parts.append(processed_content["data_insights"][:1000] + "...")
        
        # Add processing statistics
        total_files = len(processed_content.get("file_metadata", []))
        successful_ocr = sum(1 for f in processed_content.get("file_metadata", []) 
                           if f.get("file_type") == "image" and f.get("ocr_success"))
        total_images = sum(1 for f in processed_content.get("file_metadata", []) 
                          if f.get("file_type") == "image")
        
        context_parts.append(f"\n\nPROCESSING SUMMARY:")
        context_parts.append(f"• Total files processed: {total_files}")
        context_parts.append(f"• Total content analyzed: {processed_content['total_content_length']} characters")
        context_parts.append(f"• Images with successful OCR: {successful_ocr}/{total_images}")
        
        # Add file type breakdown
        file_types = {}
        for file_info in processed_content.get("file_metadata", []):
            file_type = file_info.get("file_type", "unknown")
            file_types[file_type] = file_types.get(file_type, 0) + 1
        
        if file_types:
            context_parts.append(f"• File type breakdown: {', '.join([f'{count} {type}' for type, count in file_types.items()])}")
        
        return "\n".join(context_parts)

    def _clean_ocr_text(self, text: str) -> str:
        """Clean and improve OCR extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = ' '.join(text.split())
        
        # Remove common OCR artifacts
        text = text.replace('|', 'I')  # Common OCR mistake
        text = text.replace('0', 'O')  # Common OCR mistake in some fonts
        text = text.replace('1', 'l')  # Common OCR mistake in some fonts
        
        # Remove lines with only special characters
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            # Keep lines that have meaningful content
            if any(c.isalnum() for c in line) and len(line.strip()) > 2:
                cleaned_lines.append(line.strip())
        
        return '\n'.join(cleaned_lines)

    def _preprocess_image_for_ocr(self, image: Image.Image) -> Image.Image:
        """Preprocess image to improve OCR accuracy"""
        try:
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Resize if image is too small (OCR works better with larger images)
            min_size = 800
            if image.width < min_size or image.height < min_size:
                # Calculate new size maintaining aspect ratio
                ratio = max(min_size / image.width, min_size / image.height)
                new_width = int(image.width * ratio)
                new_height = int(image.height * ratio)
                image = image.resize((new_width, new_height), Image.Resampling.LANCZOS)
            
            # Convert to grayscale for better OCR (optional, but often improves accuracy)
            # Uncomment the next line if you want grayscale processing
            # image = image.convert('L')
            
            return image
            
        except Exception as e:
            self.logger.debug(f"Image preprocessing failed: {e}")
            return image  # Return original image if preprocessing fails
